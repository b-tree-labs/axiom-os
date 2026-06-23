# Copyright (c) 2026 The University of Texas at Austin
# Copyright (c) 2026 B-Tree Labs
# SPDX-License-Identifier: Apache-2.0

"""``axi pub`` — document lifecycle management (PRESS).

## Canonical vocabulary (v0.30)

Action verbs — produce or move artifacts:

    axi pub draft <file>          Render a draft locally (no upload)
    axi pub publish <file>        Draft + publish to configured storage
    axi pub publish --draft       Publish in draft state (review period)
    axi pub publish --all         Batch publish changed docs
    axi pub review [file]         Interactive human-in-the-loop review
    axi pub review --quick        Fast one-shot approve/reject
    axi pub pull [<doc_id>]       Pull external doc → update local .md
    axi pub pull --all            Pull all tracked docs
    axi pub scan                  Scan source folders vs manifests
    axi pub onboard <id> <file>   Register a doc into the manifest
    axi pub watch                 Watch source dirs; auto-publish on save
    axi pub push <path>           Storage-only upload (multi-section aware)
    axi pub assemble <manifest>   Assemble multi-section doc from .compile.yaml
    axi pub check-links           Verify cross-doc links resolve

Query verbs — show state, no side effects:

    axi pub overview              Dashboard of the document ecosystem
    axi pub status [file]         Per-doc status (or full set when omitted)
    axi pub diff                  Show docs changed since last publish
    axi pub providers             List available storage / generation providers
    axi pub standards             List PRESS named standards (ADR-058)

Standards execution (ADR-058):

    axi pub do <name>             Run a named PRESS standard bundle

## Deprecated aliases (removed in v0.31)

    axi pub generate <file>       → use `axi pub draft <file>`

Each alias emits a stderr deprecation notice and dispatches to its
canonical handler.
"""

from __future__ import annotations

import argparse
import logging
import sys
from pathlib import Path

from axiom.infra.cli_format import section_header, separator
from axiom.infra.time_utils import time_ago

logger = logging.getLogger(__name__)


def _all_source_docs(engine) -> list[str]:
    """Collect all docs from configured source_dirs."""
    from axiom import REPO_ROOT
    from axiom.infra.config_loader import load_yaml

    config_path = REPO_ROOT / ".neut" / "publisher" / "workflow.yaml"
    cfg = load_yaml(config_path)
    folders = cfg.get("source_dirs", cfg.get("folders", []))
    if not folders:
        folders = [
            {"path": "docs/prds", "pattern": "*.md"},
            {"path": "docs/adrs", "pattern": "*.md"},
            {"path": "docs/specs", "pattern": "*.md"},
        ]
    docs = []
    for folder_cfg in folders:
        folder = REPO_ROOT / folder_cfg["path"]
        pattern = folder_cfg.get("pattern", "*.md")
        if folder.exists():
            for md_file in sorted(folder.glob(pattern)):
                if not md_file.name.startswith("_") and md_file.name != "README.md":
                    docs.append(str(md_file.relative_to(REPO_ROOT)))
    return docs


def cmd_publish(args: argparse.Namespace) -> None:
    """Publish document(s) to configured storage."""
    from .engine import PublisherEngine

    engine = PublisherEngine()
    force = getattr(args, 'force', False)

    if args.all and force:
        # Force-publish everything (explicit opt-in)
        changed = engine.diff()
        if not changed:
            changed = _all_source_docs(engine)
        for doc in changed:
            source = engine.config.repo_root / doc
            if source.exists():
                engine.publish(source, storage_override=args.endpoint,
                               draft=args.draft, force=force)
    elif args.all or (not args.file and not getattr(args, 'path', None)):
        # Default behavior: publish only what changed (incremental)
        changed = engine.diff()
        if not changed:
            print("  No documents changed since last publish.")
            print("  Use --force to republish everything.")
            return
        print(f"\n  {len(changed)} document(s) changed:\n")
        for doc in changed:
            print(f"    {doc}")
        print()
        for doc in changed:
            source = engine.config.repo_root / doc
            if source.exists():
                engine.publish(
                    source,
                    storage_override=args.endpoint,
                    draft=args.draft,
                    force=force,
                )
    elif args.file:
        source = Path(args.file).resolve()
        if not source.exists():
            print(f"✗ File not found: {args.file}")
            sys.exit(1)
        result = engine.publish(
            source,
            storage_override=args.endpoint,
            draft=args.draft,
            force=force,
        )

        # Show publication details if available
        if result and isinstance(result, dict):
            print("\n" + section_header("✓ Document Published Successfully", width=70))
            if result.get('version'):
                print(f"Version:  {result['version']}")
            if result.get('storage'):
                print(f"Storage:  {result['storage']}")
            if result.get('url'):
                print(f"URL:      {result['url']}")
            if result.get('git_footer'):
                print(f"Git URL:  {result['git_footer']}")
            print(separator("=", width=70) + "\n")
    else:
        print("Specify a file or use --all --changed-only")
        sys.exit(1)


def cmd_draft(args: argparse.Namespace) -> None:
    """Render a draft artifact locally without uploading.

    Renamed from ``cmd_generate`` 2026-06-01 to give writers the verb
    pair they actually use (``draft`` → ``publish``). ``generate``
    survives as a deprecated alias through v0.30.x.
    """
    from .engine import PublisherEngine

    source = Path(args.file).resolve()
    if not source.exists():
        print(f"File not found: {args.file}")
        sys.exit(1)

    engine = PublisherEngine()
    output = engine.generate(source)
    print(f"\nGenerated: {output}")



def cmd_status(args: argparse.Namespace) -> None:
    """Show document status with detailed information."""
    from .engine import PublisherEngine

    engine = PublisherEngine()

    # Load manifests to check for SharePoint URLs
    manifests = engine.load_manifests([
        engine.config.repo_root / "docs" / "prd",
        engine.config.repo_root / "docs" / "specs",
    ])

    if args.file:
        source = Path(args.file).resolve()
        docs = engine.status(source)
    else:
        docs = engine.status()

    if not docs:
        print("No tracked documents.")
        return

    # Detailed view
    for _, doc in enumerate(docs):
        print("\n" + separator())
        print(f"📄 {doc.doc_id}")
        print(separator())

        # Status badge
        status_icon = "✓" if doc.status == "published" else "○" if doc.status == "draft" else "?"
        print(f"  Status:    {status_icon} {doc.status.upper()}")

        # Source file
        source_path = engine.config.repo_root / doc.source_path
        file_exists = source_path.exists()
        file_icon = "✓" if file_exists else "✗"
        print(f"  Source:    {file_icon} {doc.source_path}")

        # Published version
        if doc.published:
            print(f"  Published: v{doc.published.version}")
            pub_ago = time_ago(doc.published.published_at)
            print(f"    • Version: {doc.published.version}")
            print(f"    • Published: {doc.published.published_at[:10]} ({pub_ago})")
            print(f"    • Provider: {doc.published.storage_provider}")
            print(f"    • URL: {doc.published.url}")

        # Draft version
        if doc.active_draft:
            draft_ago = time_ago(doc.active_draft.published_at)
            print(f"  Draft:     v{doc.active_draft.version}")
            print(f"    • Version: {doc.active_draft.version}")
            print(f"    • Published: {doc.active_draft.published_at[:10]} ({draft_ago})")

        # Manifest status
        in_manifest = doc.doc_id in manifests
        manifest_icon = "✓" if in_manifest else "○"
        if in_manifest:
            manifest_entry = manifests[doc.doc_id]
            sharepoint_url = manifest_entry.get("published_url")
            sharepoint_icon = "✓" if sharepoint_url else "✗"
            print(f"  Manifest:  {manifest_icon} In manifest")
            if sharepoint_url:
                print(f"  SharePoint: {sharepoint_icon} Configured")
            else:
                print(f"  SharePoint: {sharepoint_icon} NOT configured")
        else:
            print(f"  Manifest:  {manifest_icon} Not in manifest (orphaned in registry)")

        # Git info
        if doc.published and doc.published.commit_sha:
            commit = doc.published.commit_sha[:8]
            print(f"  Git:       • Commit: {commit}")
            if doc.published.generation_provider:
                print("             • Footer: Included (git source URL, version, date)")

    print("\n" + separator())
    print(f"Total: {len(docs)} document(s)")
    in_manifest_count = sum(1 for d in docs if d.doc_id in manifests)
    print(f"In manifest: {in_manifest_count}/{len(docs)}")
    print(separator() + "\n")


def cmd_check_links(args: argparse.Namespace) -> None:
    """Verify cross-document links."""
    from .engine import PublisherEngine

    engine = PublisherEngine()
    results = engine.check_links()

    valid = results.get("valid", [])
    missing = results.get("missing", [])

    print("\n" + section_header("Cross-Document Link Verification"))

    if not valid and not missing:
        print("\nℹ No documents in registry. Publish some documents first.")
        print("  Run: axi pub publish <file>")
        print(separator("=") + "\n")
        return

    if valid:
        print(f"\n✓ Valid Links ({len(valid)}):")
        print(separator())
        for path in sorted(valid):
            doc_id = Path(path).stem
            print(f"  ✓ {doc_id:<40} {path}")

    if missing:
        print(f"\n✗ Missing Source Files ({len(missing)}):")
        print(separator())
        for path in sorted(missing):
            doc_id = Path(path).stem
            print(f"  ✗ {doc_id:<40} {path}")
        print("\n⚠ These documents are in the registry but source files not found.")
        print("  Either restore the files or remove from registry.")

    total = len(valid) + len(missing)
    if total > 0:
        pct = (len(valid) / total) * 100 if total > 0 else 0
        print("\n" + separator())
        print(f"Link Health: {len(valid)}/{total} ({pct:.0f}%)")
        if pct == 100:
            print("✓ All links are valid!")
        elif pct >= 80:
            print("⚠ Most links are valid, but some sources are missing.")
        else:
            print("✗ Many links are broken. Run 'axi pub scan' for details.")

    print(separator("=") + "\n")


def cmd_diff(args: argparse.Namespace) -> None:
    """Show documents changed since last publish."""
    from .engine import PublisherEngine

    engine = PublisherEngine()
    changed = engine.diff()

    print("\n" + section_header("Changed Documents Since Last Publish"))

    if not changed:
        print("\n✓ No documents changed since last publish.")
        print("\n  To publish changed docs: axi pub publish --all --changed-only")
    else:
        print(f"\n⚠ {len(changed)} file(s) changed:")
        print(separator())
        for doc in sorted(changed):
            doc_id = Path(doc).stem
            print(f"  • {doc_id:<40} {doc}")

        print("\n" + separator())
        print("To publish these changes:")
        print("  axi pub publish --all --changed-only")

    print(separator("=") + "\n")


def cmd_pull(args: argparse.Namespace) -> None:
    """Pull document(s) from external storage → update local .md files.

    This is the reverse of `publish`:
    - Fetches the latest version from O365/OneDrive
    - Extracts text, comments, tracked changes
    - Updates the local .md file (or shows diff if --dry-run)
    """
    from .engine import PublisherEngine

    engine = PublisherEngine()

    if args.all:
        # Pull all tracked documents
        docs = engine.status()
        tracked = [d for d in docs if d.published and d.published.storage_id]
        if not tracked:
            print("No documents with external storage tracking.")
            return

        print(f"Pulling {len(tracked)} document(s) from external storage...\n")
        for doc in tracked:
            print(f"  {doc.doc_id}...")
            try:
                result = engine.pull(
                    doc.doc_id,
                    dry_run=args.dry_run,
                    include_comments=args.comments,
                )
                if result.get("changed"):
                    if args.dry_run:
                        print("    → Would update (diff available)")
                    else:
                        print(f"    → Updated {result.get('source_path')}")
                else:
                    print("    → No changes")
            except Exception as e:
                print(f"    ✗ Error: {e}")
    elif args.doc_id:
        # Pull specific document
        result = engine.pull(
            args.doc_id,
            dry_run=args.dry_run,
            include_comments=args.comments,
        )
        if args.dry_run:
            if result.get("diff"):
                print("Changes detected:\n")
                print(result["diff"])
            else:
                print("No changes detected.")
        else:
            if result.get("changed"):
                print(f"Updated: {result.get('source_path')}")
                if result.get("comments"):
                    print(f"  {len(result['comments'])} comment(s) extracted")
            else:
                print("No changes detected.")
    else:
        print("Specify a doc_id or use --all")
        sys.exit(1)


def cmd_providers(args: argparse.Namespace) -> None:
    """List all available providers with descriptions."""
    # Ensure providers are registered
    try:
        import axiom.extensions.builtins.publishing.providers  # noqa: F401
    except ImportError:
        pass

    from .config import load_config
    from .factory import PublisherFactory

    config = load_config()
    all_providers: dict[str, list[str]] = PublisherFactory.available()  # type: ignore[assignment]

    print("\n" + section_header("Publisher Providers"))

    category_info = {
        "generation": {
            "label": "Generation (Markdown → Artifact)",
            "description": "Converts .md source files to publishable formats (DOCX, PDF, etc.)",
        },
        "storage": {
            "label": "Storage (Upload & URLs)",
            "description": "Manages artifact storage and provides canonical URLs",
        },
        "feedback": {
            "label": "Feedback (Reviewer Comments)",
            "description": "Extracts and tracks reviewer feedback from published artifacts",
        },
        "notification": {
            "label": "Notifications (Alerts)",
            "description": "Sends notifications when documents are published",
        },
        "embedding": {
            "label": "Embedding (RAG Indexing)",
            "description": "Indexes documents for semantic search and RAG applications",
        },
    }

    active_map = {
        "generation": config.generation.provider,
        "storage": config.storage.provider,
        "feedback": config.feedback.provider,
        "notification": config.notification.provider,
        "embedding": config.embedding.provider if config.embedding_enabled else None,
    }

    for category, names in all_providers.items():
        info = category_info.get(category, {})
        label = info.get("label", category)
        desc = info.get("description", "")
        active = active_map.get(category, "")

        print(f"\n{label}")
        if desc:
            print(f"  {desc}")
        print(separator(width=78))

        if names:
            for name in sorted(names):
                marker = " ★" if name == active else ""
                print(f"  {name:<30}{marker}")
        else:
            print("  (no providers registered)")

    print("\n" + separator())
    print("★ = currently active (from .publisher.yaml)")
    print(separator("=") + "\n")


def cmd_watch(args: argparse.Namespace) -> None:
    """Watch source directories and auto-publish on save."""
    from .watcher import PublishWatcher

    cooldown = args.cooldown
    if not cooldown:
        try:
            from axiom.extensions.builtins.settings.store import SettingsStore
            cooldown = int(SettingsStore().get("publisher.cooldown_seconds", 300))
        except Exception:
            cooldown = 300

    watcher = PublishWatcher(
        poll_interval=args.interval,
        cooldown=cooldown,
    )
    watcher.run()


def cmd_scan(args: argparse.Namespace) -> None:
    """Scan folders for markdown files and compare against manifests."""
    from .engine import PublisherEngine

    engine = PublisherEngine()

    # Default to docs/prds and docs/specs if no folders specified
    if args.folders:
        folders = [Path(f).resolve() for f in args.folders]
    else:
        repo = engine.config.repo_root
        folders = [
            repo / "docs" / "prd",
            repo / "docs" / "specs",
        ]

    results = engine.scan_docs(folders)
    manifests = engine.load_manifests(folders)

    print("\n" + section_header("Document Manifest Scan"))

    tracked = results.get("tracked", [])
    untracked = results.get("untracked", [])
    orphaned = results.get("orphaned", [])

    # Tracked documents with SharePoint status
    if tracked:
        print(f"\n✓ Tracked in Manifest ({len(tracked)}):")
        print(separator())

        # Group by manifest
        by_manifest = {}
        for doc_id in sorted(tracked):
            if doc_id in manifests:
                manifest_path = manifests[doc_id]["manifest_path"]
                manifest_name = manifest_path.parent.name
                if manifest_name not in by_manifest:
                    by_manifest[manifest_name] = []
                by_manifest[manifest_name].append((doc_id, manifests[doc_id]))

        for manifest_name in sorted(by_manifest.keys()):
            print(f"\n  📁 {manifest_name}/")
            for doc_id, entry in by_manifest[manifest_name]:
                has_url = entry.get("published_url") is not None
                url_icon = "🔗" if has_url else "○"
                url_status = "SharePoint configured" if has_url else "No SharePoint URL"
                print(f"    {url_icon} {doc_id:<40} [{url_status}]")

    # Untracked files
    if untracked:
        print(f"\n⚠ Untracked Files ({len(untracked)}):")
        print(separator())
        for filename in sorted(untracked):
            print(f"  • {filename}")
        print("\n  To add: axi pub onboard <doc_id> <filename> --folder <path>")

    # Orphaned entries
    if orphaned:
        print(f"\n✗ Orphaned Manifest Entries ({len(orphaned)}):")
        print(separator())
        for doc_id in sorted(orphaned):
            if doc_id in manifests:
                source_path = manifests[doc_id]["source_path"]
                print(f"  • {doc_id} (expected: {source_path})")
        print("\n  These are tracked in manifests but files don't exist on disk.")
        print("  Either: 1) Restore the file, or 2) Remove from manifest")

    # Summary
    print("\n" + separator())
    total_docs = len(tracked) + len(untracked)
    total_configured = sum(
        1 for doc_id in tracked
        if doc_id in manifests and manifests[doc_id].get("published_url")
    )
    print("Summary:")
    print(f"  Total documents found: {total_docs}")
    print(f"  In manifest: {len(tracked)}")
    print(f"  With SharePoint URL: {total_configured}/{len(tracked)}")

    if not untracked and not orphaned:
        print("\n  ✓ All documents tracked and accounted for!")

    print(separator("=") + "\n")


def cmd_onboard(args: argparse.Namespace) -> None:
    """Add a document to a manifest."""
    from .engine import PublisherEngine

    if not args.doc_id or not args.file:
        print("Usage: axi pub onboard <doc_id> <file> [--folder <path>] [--url <sharepoint_url>]")
        sys.exit(1)

    engine = PublisherEngine()

    # Determine manifest folder
    if args.folder:
        manifest_dir = Path(args.folder).resolve()
    else:
        # Infer from file path or default to docs/prds
        file_path = Path(args.file).resolve()
        if "specs" in str(file_path):
            manifest_dir = engine.config.repo_root / "docs" / "specs"
        else:
            manifest_dir = engine.config.repo_root / "docs" / "prds"

    manifest_path = manifest_dir / ".publisher.json"
    source_path = Path(args.file)

    # Verify source file exists
    if not source_path.exists():
        source_path = manifest_dir / args.file
        if not source_path.exists():
            print(f"✗ File not found: {args.file}")
            sys.exit(1)

    success = engine.onboard_doc(
        doc_id=args.doc_id,
        source_path=source_path,
        manifest_path=manifest_path,
        published_url=args.url or None,
    )

    if success:
        print("\n" + section_header("✓ Document Onboarded Successfully", width=70))
        print(f"Manifest:  {manifest_path.relative_to(engine.config.repo_root)}")
        print(f"Doc ID:    {args.doc_id}")
        print(f"Source:    {source_path.relative_to(engine.config.repo_root)}")
        if args.url:
            print(f"Published: {args.url}")
        else:
            print("Published: (not set)")
            print("\n💡 Tip: Set SharePoint URL with: axi pub onboard <id> <file> --url <url>")
        print(separator("=", width=70) + "\n")
    else:
        print("✗ Failed to onboard document")
        sys.exit(1)


def cmd_review(args: argparse.Namespace) -> None:
    """Interactive human-in-the-loop review of a draft document."""
    from axiom.review.adapters.draft_adapter import (
        DraftReviewAdapter,
        create_draft_session,
        find_draft,
    )
    from axiom.review.models import ReviewSessionStore
    from axiom.review.runner import ReviewRunner

    store = ReviewSessionStore()

    # --status: show active review sessions without entering review
    if getattr(args, "review_status", False):
        sessions = store.list_active()
        if not sessions:
            print("No active review sessions.")
            return
        print(f"\nActive review sessions: {len(sessions)}\n")
        for s in sessions:
            reviewed, total = s.progress
            print(f"  {s.session_id}")
            print(f"    Source:   {s.source}")
            print(f"    Progress: {reviewed}/{total} reviewed")
            print()
        return

    # Find the draft to review
    file_arg = getattr(args, "file", None)
    draft_path = find_draft(file_arg=file_arg)
    if not draft_path:
        if file_arg:
            print(f"Draft not found: {file_arg}")
        else:
            print("No draft files found in tools/agents/drafts/")
            print("Generate one with: axi signal draft")
        sys.exit(1)

    print(f"File: {draft_path.name}")

    # --chat: hand off to conversational review in axi chat
    if getattr(args, "chat", False):
        from axiom.extensions.builtins.chat.entry import enter_chat

        content = draft_path.read_text(encoding="utf-8")
        enter_chat(
            context_markdown=(
                f"# Review: {draft_path.name}\n\n"
                f"The user wants to review this draft conversationally.\n"
                f"Use the review_start, review_get_item, review_decide, "
                f"review_progress, and review_complete tools to drive the review.\n\n"
                f"---\n\n{content}"
            ),
            title=f"Review: {draft_path.stem}",
            suggestions=[
                "Let's review this draft",
                "Show me the first section",
                "Accept everything and finalize",
            ],
            source="neut_doc_review",
        )
        return

    fresh = getattr(args, "fresh", False)
    quick = getattr(args, "quick", False)
    session = create_draft_session(draft_path, store, fresh=fresh)

    adapter = DraftReviewAdapter()
    runner = ReviewRunner(adapter, store)
    runner.run(session, quick=quick)


def cmd_overview(args: argparse.Namespace) -> None:
    """Show dashboard overview of document ecosystem."""
    from .engine import PublisherEngine

    engine = PublisherEngine()

    # Load all manifests
    manifests = engine.load_manifests([
        engine.config.repo_root / "docs" / "prd",
        engine.config.repo_root / "docs" / "specs",
    ])

    # Scan documents
    prd_docs = engine.scan_docs([engine.config.repo_root / "docs" / "prd"])
    spec_docs = engine.scan_docs([engine.config.repo_root / "docs" / "specs"])

    total_tracked = len(prd_docs.get("tracked", [])) + len(spec_docs.get("tracked", []))
    total_untracked = len(prd_docs.get("untracked", [])) + len(spec_docs.get("untracked", []))

    # Count documents with published URLs
    with_urls = 0
    without_urls = 0
    for _doc_id, entry in manifests.items():
        if entry.get("published_url"):
            with_urls += 1
        else:
            without_urls += 1

    # Check link health
    links = engine.check_links()
    valid_links = len(links.get("valid", []))
    broken_links = len(links.get("missing", []))

    print("\n" + section_header("Publisher Ecosystem Overview"))

    print("\n📋 DOCUMENTS")
    print(separator())
    print(f"  Total Tracked:           {total_tracked}")
    print(f"  Untracked Files:         {total_untracked}")
    print(f"  With SharePoint URLs:    {with_urls}")
    print(f"  Awaiting Configuration:  {without_urls}")

    print("\n🔗 LINKS")
    print(separator())
    print(f"  Valid Links:             {valid_links}")
    print(f"  Broken/Missing:          {broken_links}")
    if valid_links + broken_links > 0:
        health_pct = (valid_links / (valid_links + broken_links)) * 100
        status = "✓ Healthy" if health_pct >= 95 else "⚠ Needs attention" if health_pct >= 80 else "✗ Critical"
        print(f"  Health:                  {health_pct:.0f}% {status}")

    print("\n📁 BY FOLDER")
    print(separator())
    print("  PRDs:")
    print(f"    • Tracked:    {len(prd_docs.get('tracked', []))}")
    print(f"    • Untracked:  {len(prd_docs.get('untracked', []))}")

    print("  Specs:")
    print(f"    • Tracked:    {len(spec_docs.get('tracked', []))}")
    print(f"    • Untracked:  {len(spec_docs.get('untracked', []))}")

    print("\n💡 NEXT STEPS")
    print(separator())
    if total_untracked > 0:
        print(f"  • {total_untracked} untracked documents — run 'axi pub scan' to see them")
    if without_urls > 0:
        print(f"  • {without_urls} documents missing SharePoint URLs — run 'axi pub onboard' to configure")
    if broken_links > 0:
        print(f"  • {broken_links} broken links — run 'axi pub check-links' to diagnose")
    if not total_untracked and not without_urls and not broken_links:
        print("  ✓ Everything configured and healthy! Ready to publish.")

    print(separator("=") + "\n")


# ---------------------------------------------------------------------------
# Command Registry (auto-synced to chat slash commands)
# ---------------------------------------------------------------------------


def cmd_standards(args: argparse.Namespace) -> None:
    """List registered PRESS standards bundles (ADR-058)."""
    from .skills import bind_default

    reg = bind_default()
    params = {}
    if getattr(args, "category", None):
        params["category"] = args.category
    result = reg.invoke("press.standards", params, None)
    if not result.ok:
        for e in result.errors:
            print(f"ERROR: {e}", file=sys.stderr)
        sys.exit(result.exit_code if hasattr(result, "exit_code") else 1)
    items = result.value.get("items") or []
    if not items:
        print("(no standards registered for that filter)")
        return
    print(f"{len(items)} PRESS standard(s):")
    for it in items:
        print(f"  {it['name']:<24} {it['description']}")
        print(f"    steps: {' → '.join(it['skills'])}")


def cmd_do_standard(args: argparse.Namespace) -> None:
    """Execute a named PRESS standard bundle (ADR-058)."""
    from .skills import bind_default

    reg = bind_default()
    params = {"name": args.name}
    if getattr(args, "source", None):
        params["source"] = args.source
    if getattr(args, "target", None):
        params["target"] = args.target
    result = reg.invoke("press.do_standard", params, None)
    if not result.ok:
        for e in result.errors:
            print(f"ERROR: {e}", file=sys.stderr)
        sys.exit(1)
    v = result.value or {}
    print(f"✓ ran standard {v.get('standard')!r} ({len(v.get('steps') or [])} step(s))")
    for step in v.get("steps", []):
        marker = "✓" if step["ok"] else "✗"
        print(f"  {marker} {step['skill']}")


COMMANDS: dict[str, str] = {
    "overview": "Dashboard of document ecosystem",
    "publish": "Draft + publish to storage",
    "draft": "Render a draft locally (no upload)",
    "review": "Interactive human-in-the-loop review",
    "pull": "Pull external doc → update .md",
    "status": "Show document status",
    "check-links": "Verify cross-doc links",
    "diff": "Show changed docs since last publish",
    "scan": "Scan folders for docs vs manifests",
    "onboard": "Add document to manifest",
    "watch": "Watch source dirs, auto-publish on save",
    "providers": "List available providers",
    "push": "Push document to storage (auto-assembles multi-section if .compile.yaml found)",
    "assemble": "Assemble a multi-section doc from .compile.yaml (plumbing)",
}


# ---------------------------------------------------------------------------
# Canonical verb dispatch — single source of truth for `axi pub <verb>`
# ---------------------------------------------------------------------------

# Maps deprecated verb → canonical replacement. Each alias emits a stderr
# notice on use (per the verb-migration discipline memory) and resolves
# to the canonical handler. Removed in v0.31.
DEPRECATED_VERB_ALIASES: dict[str, str] = {
    "generate": "draft",
}


def resolve_verb_handler(verb: str):
    """Return the handler callable for ``verb``, following alias links."""
    canonical = DEPRECATED_VERB_ALIASES.get(verb, verb)
    return _verb_handlers().get(canonical)


def _verb_handlers() -> dict[str, "callable"]:
    """Map canonical verb → handler function. Lazy so the cmd_* names
    resolve at call time regardless of definition order."""
    g = globals()
    return {
        "overview":    g["cmd_overview"],
        "publish":     g["cmd_publish"],
        "draft":       g["cmd_draft"],
        "review":      g["cmd_review"],
        "pull":        g["cmd_pull"],
        "status":      g["cmd_status"],
        "check-links": g["cmd_check_links"],
        "diff":        g["cmd_diff"],
        "scan":        g["cmd_scan"],
        "onboard":     g["cmd_onboard"],
        "watch":       g["cmd_watch"],
        "providers":   g["cmd_providers"],
        "push":        g["cmd_push"],
        "assemble":    g["cmd_assemble"],
        "standards":   g["cmd_standards"],
        "do":          g["cmd_do_standard"],
    }


def emit_deprecation_warning_if_alias(verb: str) -> None:
    """If ``verb`` is a deprecated alias, print a one-line stderr notice."""
    if verb in DEPRECATED_VERB_ALIASES:
        canonical = DEPRECATED_VERB_ALIASES[verb]
        print(
            f"axi pub {verb}: deprecated; use `axi pub {canonical}`. "
            f"Alias removed in v0.31.",
            file=sys.stderr,
        )


def build_argparse(parser: argparse.ArgumentParser) -> argparse.ArgumentParser:
    """Populate ``parser`` with the canonical ``axi pub`` subverbs.

    Wraps :func:`get_parser` so callers that already own a parent parser
    (e.g. the ``axi`` top-level dispatcher and the vocabulary tests) can
    reuse the same verb-set definition.
    """
    # Build the subparser tree on the supplied parser. We do this by
    # delegating to the existing ``get_parser`` body — which constructs
    # its own parser — and copying the subparser actions across. Less
    # surgery than refactoring ``get_parser`` itself.
    src = get_parser()
    # Walk src's _actions for the subparser action and re-register on
    # the target parser. Simpler: just transplant.
    src_sub = next(
        a for a in src._actions
        if isinstance(a, argparse._SubParsersAction)
    )
    dst_sub = parser.add_subparsers(dest="command")
    # Build a {choice → help} map from the source's _choices_actions so
    # we can preserve help text when transplanting choices below.
    src_help_by_name = {
        a.dest: a.help for a in src_sub._choices_actions
    }
    # Internal name argparse uses for the pseudo-action class differs
    # by Python version. ``_SubParsersAction._ChoicesPseudoAction`` is
    # the canonical access path.
    PseudoAction = type(src_sub._choices_actions[0]) if src_sub._choices_actions else None
    for name, sub in src_sub.choices.items():
        dst_sub._name_parser_map[name] = sub
        if PseudoAction is not None:
            action = PseudoAction(name, [], src_help_by_name.get(name))
            dst_sub._choices_actions.append(action)
    return parser


def get_parser() -> argparse.ArgumentParser:
    """Build and return the argument parser.

    Exposed for CLI registry introspection. Commands auto-sync to chat.
    """
    parser = argparse.ArgumentParser(
        prog="axi pub",
        description="Document lifecycle management",
    )

    subparsers = parser.add_subparsers(dest="command")

    # overview
    subparsers.add_parser("overview", help="Dashboard of document ecosystem")

    # publish
    pub_parser = subparsers.add_parser("publish", help="Generate + publish to storage")
    pub_parser.add_argument("file", nargs="?", help="Markdown file to publish")
    pub_parser.add_argument("--draft", action="store_true", help="Publish as draft")
    pub_parser.add_argument("--all", action="store_true", help="Batch publish")
    pub_parser.add_argument("--changed-only", action="store_true", help="Only changed docs")
    pub_parser.add_argument("--endpoint", help="Override storage provider (e.g., 'local')")
    pub_parser.add_argument("--force", action="store_true", help="Force publish even if no changes detected")

    # review
    rev_parser = subparsers.add_parser("review", help="Interactive human-in-the-loop review")
    rev_parser.add_argument("file", nargs="?", help="Draft file to review (default: most recent)")
    rev_parser.add_argument("--fresh", action="store_true", help="Start over, discard previous progress")
    rev_parser.add_argument("--quick", action="store_true", help="Fast one-shot review (approve/reject all)")
    rev_parser.add_argument("--status", action="store_true", dest="review_status", help="Show active review sessions")
    rev_parser.add_argument("--chat", action="store_true", help="Review conversationally via axi chat")

    # draft — new in v0.30; pairs with `publish` for the writer's vocabulary
    draft_parser = subparsers.add_parser(
        "draft",
        help="Render a draft locally (no upload)",
    )
    draft_parser.add_argument("file", help="Markdown source file")

    # generate — deprecated alias for `draft`; removed in v0.31
    gen_parser = subparsers.add_parser(
        "generate",
        help="(deprecated alias for `draft`; removed in v0.31)",
    )
    gen_parser.add_argument("file", help="Markdown source file")

    # pull
    pull_parser = subparsers.add_parser("pull", help="Pull external doc → update .md")
    pull_parser.add_argument("doc_id", nargs="?", help="Document ID to pull")
    pull_parser.add_argument("--all", action="store_true", help="Pull all tracked docs")
    pull_parser.add_argument("--dry-run", action="store_true", help="Show diff without updating")
    pull_parser.add_argument("--comments", action="store_true", help="Include comments in output")

    # status
    stat_parser = subparsers.add_parser("status", help="Show document status")
    stat_parser.add_argument("file", nargs="?", help="Specific file (optional)")

    # check-links
    subparsers.add_parser("check-links", help="Verify cross-doc links")

    # diff
    subparsers.add_parser("diff", help="Show changed docs since last publish")

    # scan
    scan_parser = subparsers.add_parser("scan", help="Scan folders for docs vs manifests")
    scan_parser.add_argument("folders", nargs="*", help="Folders to scan (default: docs/prds docs/specs)")

    # onboard
    onboard_parser = subparsers.add_parser("onboard", help="Add document to manifest")
    onboard_parser.add_argument("doc_id", help="Document identifier")
    onboard_parser.add_argument("file", help="Path to .md file")
    onboard_parser.add_argument("--folder", help="Manifest folder (default: inferred from file)")
    onboard_parser.add_argument("--url", help="SharePoint URL (optional)")

    # watch
    watch_parser = subparsers.add_parser("watch", help="Watch source dirs, auto-publish on save")
    watch_parser.add_argument("--interval", type=int, default=10, help="Poll interval in seconds (default: 10)")
    watch_parser.add_argument("--cooldown", type=int, default=0, help="Override cooldown seconds (default: from settings)")

    # providers
    subparsers.add_parser("providers", help="List available providers")

    # push — primary happy-path command (auto-detects .compile.yaml)
    push_parser = subparsers.add_parser(
        "push", help="Push document to storage (auto-assembles multi-section if .compile.yaml found)"
    )
    push_parser.add_argument("path", nargs="?", help="File or directory to push")
    push_parser.add_argument("--all", action="store_true", help="Push all .md files in configured folders")
    push_parser.add_argument("--draft", action="store_true", help="Publish as draft")
    push_parser.add_argument("--endpoint", help="Override storage provider (e.g. 'local', 'onedrive')")
    push_parser.add_argument("--headed", action="store_true", help="Show browser window (for first-time login)")
    push_parser.add_argument("--force", action="store_true", help="Force re-publish even if unchanged")

    # assemble — plumbing for multi-section compilation
    asm_parser = subparsers.add_parser(
        "assemble", help="Assemble a multi-section doc from .compile.yaml (plumbing)"
    )
    asm_parser.add_argument("manifest", help="Path to .compile.yaml")
    asm_parser.add_argument("--output", "-o", help="Output file path (default: <output>.assembled.md)")

    # standards — list registered PRESS standards (ADR-058)
    std_parser = subparsers.add_parser(
        "standards",
        help="List PRESS named standards bundles (ADR-058)",
    )
    std_parser.add_argument(
        "--category", help="Filter to one category (default: all)"
    )

    # do — execute a named standard
    do_parser = subparsers.add_parser(
        "do",
        help="Execute a named PRESS standard (ADR-058)",
    )
    do_parser.add_argument("name", help="Standard name (see `axi pub standards`)")
    do_parser.add_argument(
        "--source", help="Source markdown path (forwarded to all steps)"
    )
    do_parser.add_argument(
        "--target", help="Target output path (for filename-preview standards)"
    )

    return parser


def _find_compile_manifest(path: Path) -> Path | None:
    """Search *path* and its parents for a .compile.yaml manifest.

    Checks:
    1. If *path* is a directory, look for .compile.yaml inside it
    2. If *path* is a file, look for .compile.yaml in the same directory
    3. Walk up to repo root (where .git lives) stopping at each level
    """
    candidates = []
    if path.is_dir():
        candidates.append(path / ".compile.yaml")
    candidates.append(path.parent / ".compile.yaml")

    for candidate in candidates:
        if candidate.exists():
            return candidate
    return None


def _assemble_from_manifest(manifest_path: Path, output_path: Path | None = None) -> Path:
    """Concatenate source files from a .compile.yaml manifest into a single .md.

    Returns the path to the assembled file.
    """
    try:
        # .compile.yaml is YAML, not TOML
        import yaml as _yaml  # type: ignore[import]
        with open(manifest_path) as f:
            manifest = _yaml.safe_load(f)
    except ImportError:
        # Fall back to basic key:value parsing for simple manifests
        manifest = {}
        with open(manifest_path) as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith("#") and ":" in line:
                    k, _, v = line.partition(":")
                    manifest[k.strip()] = v.strip().strip('"')

    sources = manifest.get("sources", [])
    title = manifest.get("title", "")
    doc_dir = manifest_path.parent

    if not sources:
        raise ValueError(f"No sources listed in {manifest_path}")

    # Assemble
    parts = []
    if title:
        parts.append(f"# {title}\n\n")
    for src in sources:
        src_path = doc_dir / src
        if not src_path.exists():
            raise FileNotFoundError(f"Source file not found: {src_path}")
        parts.append(src_path.read_text(encoding="utf-8"))
        if not parts[-1].endswith("\n"):
            parts.append("\n")
        parts.append("\n")

    assembled = "".join(parts)

    if output_path is None:
        output_name = manifest.get("output", "assembled")
        output_path = doc_dir / f"{output_name}.assembled.md"

    output_path.write_text(assembled, encoding="utf-8")
    return output_path


def cmd_assemble(args: argparse.Namespace) -> None:
    """Assemble a multi-section document from a .compile.yaml manifest.

    This is the plumbing command. For the happy path, use `axi pub push`
    which detects and runs assembly automatically.
    """
    manifest_path = Path(args.manifest).resolve()
    if not manifest_path.exists():
        print(f"✗ Manifest not found: {args.manifest}", file=sys.stderr)
        sys.exit(1)

    output_path = Path(args.output).resolve() if args.output else None

    try:
        result = _assemble_from_manifest(manifest_path, output_path)
        print(f"✓ Assembled → {result}")
    except (FileNotFoundError, ValueError) as exc:
        print(f"✗ Assembly failed: {exc}", file=sys.stderr)
        sys.exit(1)


def cmd_push(args: argparse.Namespace) -> None:
    """Push a document (or batch) to configured storage.

    Two modes:
    - Single/directory: generates + publishes via engine (supports .compile.yaml assembly)
    - --all + --endpoint onedrive: batch generate + upload via Playwright

    First run with --headed opens a browser for Microsoft login.
    """
    import tempfile

    from .engine import PublisherEngine

    engine = PublisherEngine()
    force = getattr(args, "force", False)
    draft = getattr(args, "draft", False)
    storage = getattr(args, "endpoint", None)
    headed = getattr(args, "headed", False)
    push_all = getattr(args, "all", False)

    # ── Batch upload (--all flag) ────────────────────────────────────────
    if push_all:
        _cmd_push_batch(args, engine, draft, storage, headed, force)
        return

    # ── Single file with browser storage ──────────────────────────────
    # Route through batch path when storage is browser-based (explicit or from config)
    configured_provider = engine.config.storage.provider if hasattr(engine, "config") else None
    effective_storage = storage or configured_provider
    if args.path and effective_storage in ("onedrive", "onedrive-graph", "box-browser"):
        _cmd_push_batch(args, engine, draft, storage or effective_storage, headed, force)
        return

    # ── Single file / directory push (original behavior) ──────────────────
    if not args.path:
        print("Usage: axi pub push <path> [--all] [--endpoint onedrive] [--headed]")
        print("\nExamples:")
        print("  axi pub push docs/prds/prd_executive.md")
        print("  axi pub push --all --endpoint onedrive --headed")
        return

    target = Path(args.path).resolve()
    if not target.exists():
        print(f"✗ Path not found: {args.path}", file=sys.stderr)
        sys.exit(1)

    # Auto-detect .compile.yaml for multi-section assembly
    assembled_tmp: Path | None = None
    source = target

    manifest = _find_compile_manifest(target)
    if manifest:
        print(f"  Assembling from {manifest.relative_to(manifest.parent.parent)}...")
        try:
            tmp = tempfile.NamedTemporaryFile(
                suffix=".assembled.md", dir=manifest.parent, delete=False,
            )
            tmp.close()
            assembled_tmp = Path(tmp.name)
            source = _assemble_from_manifest(manifest, assembled_tmp)
            print(f"  Assembly complete ({source.stat().st_size // 1024}KB)")
        except (FileNotFoundError, ValueError) as exc:
            print(f"✗ Assembly failed: {exc}", file=sys.stderr)
            sys.exit(1)

    # Publish via engine
    try:
        result = engine.publish(
            source, storage_override=storage, draft=draft, force=force,
        )
        if result and isinstance(result, dict):
            print("\n" + section_header("✓ Published successfully", width=70))
            if result.get("version"):
                print(f"Version:  {result['version']}")
            if result.get("storage"):
                print(f"Storage:  {result['storage']}")
            if result.get("url"):
                print(f"URL:      {result['url']}")
            print(separator("=", width=70) + "\n")
    except Exception as e:
        print(f"✗ Publish failed: {e}", file=sys.stderr)
        raise
    finally:
        if assembled_tmp and assembled_tmp.exists():
            assembled_tmp.unlink()


def _cmd_push_batch(args, engine, draft, storage, headed, force):
    """Generate + upload via browser storage provider."""
    from axiom import REPO_ROOT

    push_all = getattr(args, "all", False)
    single_path = getattr(args, "path", None)
    config_path = REPO_ROOT / ".publisher.yaml"

    # (source_md, docx_path, onedrive_subfolder)
    files_to_push: list[tuple[Path, Path, str]] = []

    if single_path and not push_all:
        # Single file or directory
        target = Path(single_path).resolve()
        if not target.exists():
            print(f"✗ Path not found: {single_path}", file=sys.stderr)
            sys.exit(1)

        # Mirror source path: docs/prds/ → requirements/
        subfolder = _source_to_subfolder(target, REPO_ROOT)

        if target.is_dir():
            for md_file in sorted(target.glob("*.md")):
                if md_file.name.startswith("_") or md_file.name == "README.md":
                    continue
                docx_path = _docx_output_path(md_file)
                if force or _needs_regeneration(md_file, docx_path):
                    files_to_push.append((md_file, docx_path, subfolder))
        elif target.suffix == ".md":
            docx_path = _docx_output_path(target)
            if force or _needs_regeneration(target, docx_path):
                files_to_push.append((target, docx_path, subfolder))
        elif target.suffix == ".docx":
            files_to_push.append((target, target, subfolder))
    else:
        # --all: collect from configured folders
        # Check workflow.yaml and .publisher.yaml for folder config
        folders = []
        from axiom.infra.config_loader import load_yaml

        for config_name in [".neut/publisher/workflow.yaml", ".publisher.yaml"]:
            cfg = load_yaml(REPO_ROOT / config_name)
            folders = cfg.get("source_dirs", cfg.get("folders", []))
            if folders:
                break

        if not folders:
            # Scan all standard doc directories
            folders = [
                {"path": "docs/prds", "pattern": "*.md"},
            {"path": "docs/adrs", "pattern": "*.md"},
            {"path": "docs/specs", "pattern": "*.md"},
            ]

        # Collect candidate files, filtering by content hash (not git diff)
        for folder_cfg in folders:
            folder = REPO_ROOT / folder_cfg["path"]
            pattern = folder_cfg.get("pattern", "*.md")
            if folder.exists():
                for md_file in sorted(folder.glob(pattern)):
                    if md_file.name.startswith("_") or md_file.name == "README.md":
                        continue
                    docx_path = _docx_output_path(md_file)
                    # Only include files that actually need regeneration
                    if force or _needs_regeneration(md_file, docx_path):
                        subfolder = _source_to_subfolder(md_file, REPO_ROOT)
                        files_to_push.append((md_file, docx_path, subfolder))

    if not files_to_push:
        print("\n  All documents up to date. Nothing to publish.")
        return

    print(f"\n  {len(files_to_push)} update(s) ready to publish:\n")
    docx_files: list[tuple[Path, Path, str]] = []  # (source_md, docx_path, subfolder)
    for source_md, docx_path, subfolder in files_to_push:
        print(f"    Generating {docx_path.name}...", end=" ", flush=True)
        docx_path = _generate_docx(source_md)
        if docx_path is None:
            # Mermaid rendering failed — skip this file
            continue
        # Write generation hash immediately so re-runs don't regenerate unchanged docs.
        # A separate .uploaded marker is set after a successful upload.
        import hashlib as _hl
        _gen_hash = _hl.sha256(source_md.read_bytes()).hexdigest()
        docx_path.with_suffix(".docx.sha256").write_text(_gen_hash, encoding="utf-8")
        print("\u2713")
        docx_files.append((source_md, docx_path, subfolder))

    # Resolve browser storage provider
    try:
        from .providers.storage.box_browser import BoxBrowserStorageProvider
        from .providers.storage.onedrive_browser import OneDriveBrowserStorageProvider
    except ImportError:
        print("\n✗ Playwright not installed. Run:")
        print("    pip install playwright && playwright install chromium")
        sys.exit(1)

    # Read config
    from axiom.infra.branding import get_branding as _gb_od
    onedrive_root = _gb_od().product_name
    site_url = ""
    from axiom.infra.config_loader import load_yaml

    cfg = load_yaml(config_path)
    if cfg:
        storage_cfg = cfg.get("storage", {})
        from axiom.infra.branding import get_branding as _gb_od2
        onedrive_root = storage_cfg.get("onedrive_root", storage_cfg.get("onedrive_folder", _gb_od2().product_name))
        site_url = storage_cfg.get("onedrive_url", "")

    if storage == "box-browser":
        provider = BoxBrowserStorageProvider({
            "folder": onedrive_root,
            "headless": not headed,
        })
        if not provider.has_session() and not headed:
            print("\n  No saved session. Run with --headed for first-time login:")
            print("    axi pub push --all --endpoint box-browser --headed\n")
            sys.exit(1)
    elif storage == "onedrive-graph":
        from .providers.storage.onedrive_graph import OneDriveGraphStorageProvider
        provider = OneDriveGraphStorageProvider({
            "folder": onedrive_root,
        })
    elif storage == "onedrive":
        provider = OneDriveBrowserStorageProvider({
            "folder": onedrive_root,
            "site_url": site_url,
            "headless": not headed,
        })
        if not provider.has_session() and not headed:
            print("\n  No saved session. Run with --headed for first-time login:")
            print("    axi pub push --all --endpoint onedrive --headed\n")
            sys.exit(1)
    else:
        # Default to browser upload (Graph API requires Azure AD credentials)
        provider = OneDriveBrowserStorageProvider({
            "folder": onedrive_root,
            "site_url": site_url,
            "headless": not headed,
        })
        if not provider.has_session() and not headed:
            print("\n  No saved session. Run with --headed for first-time login:")
            print("    axi pub push --all --headed\n")
            sys.exit(1)

    # Build per-file folder paths
    # docx_files is (source_md, docx_path, subfolder)
    just_files = [df[1] for df in docx_files]  # docx_path
    per_file_folders = [
        f"{onedrive_root}/{df[2]}" if df[2] else onedrive_root
        for df in docx_files
    ]

    # Show target folders and confirm
    unique_folders = sorted(set(per_file_folders))
    print(f"\n  Publishing {len(just_files)} document(s) to:")
    for folder in unique_folders:
        count = per_file_folders.count(folder)
        print(f"    OneDrive/{folder}/  ({count} files)")
    print()

    try:
        response = input("  Press Enter to confirm, or type a different folder: ").strip()
        if response:
            # User wants a different folder
            per_file_folders = [response] * len(just_files)
            print(f"  → Publishing to OneDrive/{response}/\n")
    except (EOFError, KeyboardInterrupt):
        print("\n  Cancelled.\n")
        return

    try:
        results = provider.upload_batch(just_files, draft=draft, headed=headed, folders=per_file_folders)
    except TypeError:
        # Provider doesn't support per-file folders
        results = provider.upload_batch(just_files, draft=draft, headed=headed)
    # Record publications in registry
    from axiom.infra.publication_registry import PublicationRegistry
    registry = PublicationRegistry()

    success = 0
    for i, (f, result) in enumerate(zip(just_files, results)):
        icon = "✓" if result.success else "✗"
        msg = result.url if result.success else result.error
        print(f"    {icon} {f.name}  {msg}")
        if result.success:
            success += 1
            source_md, docx_path, subfolder = docx_files[i]
            # Mark this specific docx as successfully uploaded (separate from generated hash)
            docx_path.with_suffix(".docx.uploaded").write_text(result.url, encoding="utf-8")

            # Record in registry so SCAN knows the baseline
            doc_id = f.stem  # e.g., "prd-executive"
            folder = per_file_folders[i] if i < len(per_file_folders) else ""
            try:
                source_path = str(source_md.relative_to(REPO_ROOT))

                registry.record_publication(
                    doc_id=doc_id,
                    source_path=source_path or f"docs/{f.stem}.md",
                    published_name=f.name,
                    endpoint=storage or "onedrive",
                    endpoint_folder=folder,
                    endpoint_modified=result.metadata.get("modified", "") if hasattr(result, "metadata") else "",
                    endpoint_item_id=result.storage_id or "",
                    endpoint_url=result.url or "",
                )
            except Exception as e:
                logger.debug("Failed to record publication: %s", e)

    print(f"\n  {success}/{len(just_files)} published successfully.\n")


def _postprocess_docx(docx_path: Path) -> None:
    """Post-process a .docx file for quality:

    1. Tables: cell borders, full page width, auto-fit column widths
    2. Diagrams: keep with preceding heading, size to page width
    3. Headings: "keep with next" so headings don't orphan from content
    4. Bookmarks: remove pandoc's heading bookmarks
    """
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches  # noqa: F401 (Emu removed)

        doc = Document(str(docx_path))

        # --- Tables: borders + full width + auto-fit columns ---
        for table in doc.tables:
            tbl = table._tbl
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement("w:tblPr")
                tbl.insert(0, tblPr)

            # Full page width (letter = 8.5" - 1" margins each side = 6.5")
            for existing in tblPr.findall(qn("w:tblW")):
                tblPr.remove(existing)
            tblW = OxmlElement("w:tblW")
            tblW.set(qn("w:w"), "5000")
            tblW.set(qn("w:type"), "pct")  # 100% of available width
            tblPr.append(tblW)

            # Calculate column widths based on content length
            # Total available width: ~9360 twips (6.5 inches × 1440 twips/inch)
            total_twips = 9360

            # Cell borders
            for row in table.rows:
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.tcPr
                    if tcPr is None:
                        tcPr = OxmlElement("w:tcPr")
                        tc.insert(0, tcPr)

                    for existing in tcPr.findall(qn("w:tcBorders")):
                        tcPr.remove(existing)

                    borders = OxmlElement("w:tcBorders")
                    for border_name in ["top", "left", "bottom", "right"]:
                        border = OxmlElement(f"w:{border_name}")
                        border.set(qn("w:val"), "single")
                        border.set(qn("w:sz"), "4")
                        border.set(qn("w:space"), "0")
                        border.set(qn("w:color"), "AAAAAA")
                        borders.append(border)
                    tcPr.append(borders)

            # Set column widths proportional to max content length per column
            num_cols = len(table.columns)
            if num_cols > 0:
                max_lens = [0] * num_cols
                for row in table.rows:
                    for j, cell in enumerate(row.cells):
                        if j < num_cols:
                            content_len = max(len(line) for line in cell.text.split("\n")) if cell.text else 1
                            max_lens[j] = max(max_lens[j], content_len)

                # Ensure minimum width and calculate proportions
                max_lens = [max(w, 3) for w in max_lens]
                total_len = sum(max_lens)
                col_widths = [int(total_twips * w / total_len) for w in max_lens]

                # Apply widths to all cells
                for row in table.rows:
                    for j, cell in enumerate(row.cells):
                        if j < num_cols:
                            tc = cell._tc
                            tcPr = tc.tcPr
                            if tcPr is None:
                                tcPr = OxmlElement("w:tcPr")
                                tc.insert(0, tcPr)
                            for existing in tcPr.findall(qn("w:tcW")):
                                tcPr.remove(existing)
                            tcW = OxmlElement("w:tcW")
                            tcW.set(qn("w:w"), str(col_widths[j]))
                            tcW.set(qn("w:type"), "dxa")
                            tcPr.append(tcW)

        # --- Headings + images: keep coupled, prevent page breaks between them ---
        for i, para in enumerate(doc.paragraphs):
            pPr = para._p.get_or_add_pPr()

            # Headings: keep with next
            if para.style and para.style.name and "Heading" in para.style.name:
                if pPr.find(qn("w:keepNext")) is None:
                    pPr.append(OxmlElement("w:keepNext"))

            # Image paragraphs: keep with previous heading AND keep lines together
            has_image = bool(para._p.findall(".//" + qn("wp:inline")))
            if has_image:
                if pPr.find(qn("w:keepNext")) is None:
                    pPr.append(OxmlElement("w:keepNext"))
                if pPr.find(qn("w:keepLines")) is None:
                    pPr.append(OxmlElement("w:keepLines"))
                # Also set keepNext on the preceding paragraph if it's not already set
                if i > 0:
                    prev_pPr = doc.paragraphs[i-1]._p.get_or_add_pPr()
                    if prev_pPr.find(qn("w:keepNext")) is None:
                        prev_pPr.append(OxmlElement("w:keepNext"))

            # Image captions: keep with previous
            if para.style and para.style.name and "Caption" in para.style.name:
                if pPr.find(qn("w:keepLines")) is None:
                    pPr.append(OxmlElement("w:keepLines"))

        # --- Images: size to 6 inches wide (fits letter with 1.25" margins) ---
        max_width = Inches(6)
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                # Find all inline drawings that reference this image
                for drawing in doc.element.body.iter(qn("wp:inline")):
                    extent = drawing.find(qn("wp:extent"))
                    if extent is not None:
                        cx = int(extent.get("cx", "0"))
                        cy = int(extent.get("cy", "0"))
                        if cx > max_width:
                            ratio = max_width / cx
                            extent.set("cx", str(int(max_width)))
                            extent.set("cy", str(int(cy * ratio)))

        # --- Remove bookmarks ---
        body = doc.element.body
        for bookmark_start in body.findall(".//" + qn("w:bookmarkStart")):
            bookmark_start.getparent().remove(bookmark_start)
        for bookmark_end in body.findall(".//" + qn("w:bookmarkEnd")):
            bookmark_end.getparent().remove(bookmark_end)

        doc.save(str(docx_path))

    except ImportError:
        pass
    except Exception as e:
        logger.warning("Docx post-processing failed: %s", e)


def _source_to_subfolder(source: Path, repo_root: Path) -> str:
    """Derive OneDrive subfolder by mirroring the source path relative to docs/.

    docs/prds/prd-executive.md → requirements
    docs/specs/spec-model-routing.md → tech-specs
    """
    resolved = source.resolve()
    try:
        rel = resolved.parent.relative_to(repo_root)
        return str(rel)
    except ValueError:
        return ""


def _docx_output_path(md_path: Path) -> Path:
    """Compute the expected .docx output path without generating."""
    from axiom import REPO_ROOT

    try:
        rel = md_path.parent.relative_to(REPO_ROOT)
        output_dir = REPO_ROOT / ".neut" / "generated" / rel
    except ValueError:
        output_dir = REPO_ROOT / ".neut" / "generated"

    return output_dir / (md_path.stem + ".docx")


def _needs_regeneration(md_path: Path, docx_path: Path) -> bool:
    """Check if the docx needs to be regenerated based on content hash.

    Hash (.docx.sha256) is written after successful generation so that
    re-runs don't regenerate unchanged docs even if the upload failed.
    A separate .uploaded marker is set after a successful upload.
    """
    import hashlib

    if not docx_path.exists():
        return True

    hash_file = docx_path.with_suffix(".docx.sha256")

    if not hash_file.exists():
        # No hash = never successfully uploaded. Regenerate.
        return True

    source_hash = hashlib.sha256(md_path.read_bytes()).hexdigest()
    stored_hash = hash_file.read_text(encoding="utf-8").strip()
    return source_hash != stored_hash


def _generate_docx(md_path: Path) -> Path | None:
    """Generate a .docx from a .md file using pandoc. Returns path to .docx."""
    import subprocess

    from axiom import REPO_ROOT

    # Mirror source folder structure: docs/prds/ → .neut/generated/requirements/
    try:
        rel = md_path.parent.relative_to(REPO_ROOT)
        output_dir = REPO_ROOT / ".neut" / "generated" / rel
    except ValueError:
        output_dir = REPO_ROOT / ".neut" / "generated"
    output_dir.mkdir(parents=True, exist_ok=True)

    docx_name = md_path.stem + ".docx"
    output_path = output_dir / docx_name

    # Get title from first line
    title = md_path.stem.replace("_", " ").replace("-", " ").title()
    try:
        first_line = md_path.read_text(encoding="utf-8").split("\n")[0]
        if first_line.startswith("# "):
            title = first_line[2:].strip()
    except Exception:
        pass

    # Pre-process: render mermaid diagrams to PNG images
    md_content = md_path.read_text(encoding="utf-8")
    processed_path = md_path  # Default: use original

    if "```mermaid" in md_content:
        try:
            from .mermaid_renderer import render_mermaid_blocks
            mermaid_result = render_mermaid_blocks(md_content, output_dir)
            processed_path = output_dir / f"{md_path.stem}.processed.md"
            processed_path.write_text(mermaid_result.content, encoding="utf-8")

            if not mermaid_result.all_succeeded:
                print(f"\n  \u2717 {md_path.name}: {mermaid_result.failed}/{mermaid_result.total} diagram(s) failed to render:")
                for idx, first_line in mermaid_result.failures:
                    print(f"      diagram {idx}: {first_line}...")
                print("    Fix the Mermaid syntax and retry. Skipping publish for this doc.\n")
                return None
        except Exception as e:
            logger.warning("Mermaid pre-processing failed: %s", e)

    # Build pandoc command
    cmd = [
        "pandoc", str(processed_path),
        "-o", str(output_path),
        "--from", "markdown",
        "--to", "docx",
        "--toc", "--toc-depth=3",
        "--metadata", f"title={title}",
        "--resource-path", str(output_dir),  # So pandoc can find rendered images
    ]

    # Use reference doc for table styles (borders, widths) if available
    ref_doc = REPO_ROOT / ".neut" / "publisher" / "reference.docx"
    if ref_doc.exists():
        cmd.extend(["--reference-doc", str(ref_doc)])

    try:
        subprocess.run(cmd, check=True, capture_output=True)

        # Clean up processed temp file
        if processed_path != md_path and processed_path.exists():
            processed_path.unlink()

        # Post-process: cell borders, full-width tables, remove bookmarks
        _postprocess_docx(output_path)
    except subprocess.CalledProcessError as e:
        print(f"\n    Warning: pandoc failed for {md_path.name}: {e.stderr.decode()[:200]}")
    except FileNotFoundError:
        print("\n    ✗ pandoc not installed. Run: brew install pandoc")

    return output_path


def main():
    """CLI entry point for axi doc."""
    parser = get_parser()
    args = parser.parse_args()

    if args.command == "overview":
        cmd_overview(args)
    elif args.command == "publish":
        cmd_publish(args)
    elif args.command == "review":
        cmd_review(args)
    elif args.command in ("draft", "generate"):
        emit_deprecation_warning_if_alias(args.command)
        cmd_draft(args)
    elif args.command == "pull":
        cmd_pull(args)
    elif args.command == "status":
        cmd_status(args)
    elif args.command == "check-links":
        cmd_check_links(args)
    elif args.command == "diff":
        cmd_diff(args)
    elif args.command == "scan":
        cmd_scan(args)
    elif args.command == "onboard":
        cmd_onboard(args)
    elif args.command == "providers":
        cmd_providers(args)
    elif args.command == "watch":
        cmd_watch(args)
    elif args.command == "push":
        cmd_push(args)
    elif args.command == "assemble":
        cmd_assemble(args)
    elif args.command == "standards":
        cmd_standards(args)
    elif args.command == "do":
        cmd_do_standard(args)
    else:
        parser.print_help()
        sys.exit(1)


if __name__ == "__main__":
    main()
